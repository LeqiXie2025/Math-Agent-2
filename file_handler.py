import os
import tempfile
import mimetypes
from werkzeug.utils import secure_filename
from typing import Tuple, Optional
import logging

logger = logging.getLogger(__name__)

try:
    # 尝试导入OCR和文档处理库
    import fitz  # PyMuPDF
    from PIL import Image
    import pytesseract
    from docx import Document

    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.warning("OCR相关库未安装，部分文件处理功能将受限")


class FileHandler:
    """文件处理器"""

    def __init__(self, upload_folder: str):
        self.upload_folder = upload_folder
        os.makedirs(upload_folder, exist_ok=True)

        # 支持的文件类型
        self.supported_extensions = {
            'image': ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff'],
            'document': ['pdf', 'docx', 'txt'],
            'archive': ['zip', 'rar']
        }

        # OCR配置
        self.ocr_enabled = HAS_OCR

    def allowed_file(self, filename: str) -> bool:
        """检查文件类型是否允许"""
        if '.' not in filename:
            return False

        ext = filename.rsplit('.', 1)[1].lower()

        # 检查所有支持的类型
        for category, extensions in self.supported_extensions.items():
            if ext in extensions:
                return True

        return False

    def save_file(self, file) -> Tuple[bool, str, str]:
        """保存上传的文件"""
        try:
            if not file or not self.allowed_file(file.filename):
                return False, "不支持的文件类型", ""

            # 安全处理文件名
            filename = secure_filename(file.filename)

            # 添加时间戳避免重复
            import time
            timestamp = int(time.time())
            name, ext = os.path.splitext(filename)
            unique_filename = f"{name}_{timestamp}{ext}"

            filepath = os.path.join(self.upload_folder, unique_filename)

            # 保存文件
            file.save(filepath)

            logger.info(f"文件保存成功: {unique_filename} -> {filepath}")
            return True, unique_filename, filepath

        except Exception as e:
            logger.error(f"保存文件失败: {str(e)}")
            return False, str(e), ""

    def extract_text(self, filepath: str) -> str:
        """从文件中提取文本"""
        try:
            if not os.path.exists(filepath):
                return "文件不存在"

            ext = self._get_file_extension(filepath)

            # 根据文件类型选择提取方法
            if ext in ['txt']:
                return self._extract_from_txt(filepath)
            elif ext in ['pdf']:
                return self._extract_from_pdf(filepath)
            elif ext in ['docx']:
                return self._extract_from_docx(filepath)
            elif ext in self.supported_extensions['image']:
                return self._extract_from_image(filepath)
            else:
                return f"不支持的文件类型: {ext}"

        except Exception as e:
            logger.error(f"提取文本失败: {str(e)}")
            return f"文本提取失败: {str(e)}"

    def _get_file_extension(self, filepath: str) -> str:
        """获取文件扩展名"""
        return os.path.splitext(filepath)[1].lower().replace('.', '')

    def _extract_from_txt(self, filepath: str) -> str:
        """从TXT文件提取文本"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            return content
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(filepath, 'r', encoding='gbk') as f:
                content = f.read()
            return content
        except Exception as e:
            logger.error(f"读取TXT文件失败: {str(e)}")
            return f"无法读取TXT文件: {str(e)}"

    def _extract_from_pdf(self, filepath: str) -> str:
        """从PDF文件提取文本"""
        if not HAS_OCR:
            return "PDF处理需要安装PyMuPDF库，请运行: pip install pymupdf"

        try:
            text = ""
            with fitz.open(filepath) as doc:
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text()

            # 如果文本为空，尝试OCR
            if not text.strip():
                text = self._ocr_pdf(filepath)

            return text if text.strip() else "PDF文件中未检测到文本内容"

        except Exception as e:
            logger.error(f"提取PDF文本失败: {str(e)}")
            return f"PDF文本提取失败: {str(e)}"

    def _ocr_pdf(self, filepath: str) -> str:
        """OCR识别PDF"""
        if not self.ocr_enabled:
            return "OCR功能未启用"

        try:
            text = ""
            with fitz.open(filepath) as doc:
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap()

                    # 保存临时图片
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        img_path = tmp.name
                        pix.save(img_path)

                    # OCR识别
                    img = Image.open(img_path)
                    page_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    text += page_text + "\n\n"

                    # 清理临时文件
                    os.unlink(img_path)

            return text

        except Exception as e:
            logger.error(f"OCR识别失败: {str(e)}")
            return f"OCR识别失败: {str(e)}"

    def _extract_from_docx(self, filepath: str) -> str:
        """从Word文档提取文本"""
        if not HAS_OCR:
            return "Word文档处理需要安装python-docx库，请运行: pip install python-docx"

        try:
            doc = Document(filepath)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text

        except Exception as e:
            logger.error(f"提取Word文档失败: {str(e)}")
            return f"Word文档提取失败: {str(e)}"

    def _extract_from_image(self, filepath: str) -> str:
        """从图片提取文本"""
        if not self.ocr_enabled:
            return "图片OCR需要安装Pillow和pytesseract库\n请运行: pip install pillow pytesseract\n并安装Tesseract OCR: https://github.com/tesseract-ocr/tesseract"

        try:
            img = Image.open(filepath)

            # 预处理图像（提高OCR准确率）
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # OCR识别
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')

            return text if text.strip() else "图片中未检测到文本内容"

        except Exception as e:
            logger.error(f"图片OCR失败: {str(e)}")
            return f"图片OCR失败: {str(e)}\n请确保已安装Tesseract OCR并添加到系统PATH"

    def get_file_info(self, filepath: str) -> dict:
        """获取文件信息"""
        try:
            stat = os.stat(filepath)
            ext = self._get_file_extension(filepath)

            return {
                'filename': os.path.basename(filepath),
                'size': stat.st_size,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'extension': ext,
                'type': mimetypes.guess_type(filepath)[0] or 'unknown'
            }
        except Exception as e:
            logger.error(f"获取文件信息失败: {str(e)}")
            return {}

    def cleanup_old_files(self, max_age_hours: int = 24):
        """清理旧文件"""
        import time
        current_time = time.time()

        try:
            for filename in os.listdir(self.upload_folder):
                filepath = os.path.join(self.upload_folder, filename)
                if os.path.isfile(filepath):
                    file_age = current_time - os.path.getmtime(filepath)
                    if file_age > max_age_hours * 3600:  # 转换为秒
                        os.remove(filepath)
                        logger.info(f"清理旧文件: {filename}")
        except Exception as e:
            logger.error(f"清理文件失败: {str(e)}")
